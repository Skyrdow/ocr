import os
import re
import time

import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    raise ValueError("GEMINI_API_KEY not found in .env file")

genai.configure(api_key=API_KEY)


def upload_file(path, mime_type=None):
    """Uploads the given file to Gemini."""
    file = genai.upload_file(path, mime_type=mime_type)
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file


def wait_for_files_active(files):
    """Waits for the given files to be active."""
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(2)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready")


def detect_anomalies(text):
    """
    Detects potential OCR errors and anomalies in the text.
    Returns a list of issues found.
    """
    issues = []

    # Check for garbled text patterns (repeated characters, unusual symbols)
    if re.search(r"(.)\1{3,}", text):  # 4 or more repeated characters
        issues.append("Repeated characters detected (possible OCR garble)")

    # Check for missing spaces between words
    if re.search(r"[a-z][A-Z]", text):  # lowercase followed by uppercase without space
        issues.append("Missing spaces between words detected")

    # Check for unusual character combinations
    unusual_patterns = [
        r"[^\w\s]{3,}",  # 3+ consecutive non-word characters
        r"\d{10,}",  # Very long numbers (possible garble)
        r"[|@#$%^&*]{2,}",  # Multiple special characters
    ]

    for pattern in unusual_patterns:
        if re.search(pattern, text):
            issues.append(f"Unusual character pattern detected: {pattern}")

    # Check for incomplete words (very short words that might be fragments)
    words = text.split()
    short_words = [word for word in words if len(word) == 1 and word.isalpha()]
    if len(short_words) > len(words) * 0.1:  # More than 10% single letters
        issues.append(
            "High frequency of single-letter words (possible word fragmentation)"
        )

    return issues


def calculate_confidence_score(text, issues):
    """
    Calculates a confidence score based on text quality and detected issues.
    Returns a score from 0-100.
    """
    base_score = 85  # Base confidence for Gemini OCR

    # Deduct points for each type of issue
    deductions = {
        "Repeated characters detected": 15,
        "Missing spaces between words detected": 10,
        "Unusual character pattern detected": 8,
        "High frequency of single-letter words": 12,
    }

    for issue in issues:
        for issue_type, deduction in deductions.items():
            if issue_type in issue:
                base_score -= deduction
                break

    # Ensure score stays within bounds
    return max(0, min(100, base_score))


def transcribe_and_translate(pdf_path):
    """
    Uploads a PDF, transcribes it, and translates it to Spanish using Gemini.
    """
    # model = genai.GenerativeModel(model_name="gemini-3-pro-preview") # Use this if you have quota
    model = genai.GenerativeModel(
        model_name="gemini-flash-latest"
    )  # Fallback to reliable model

    # Upload the file
    pdf_file = upload_file(pdf_path, mime_type="application/pdf")

    # Wait for processing
    wait_for_files_active([pdf_file])

    # Prompt for transcription and translation with enhanced quality assessment
    prompt = """
    Please perform the following tasks for the attached PDF file:
    1. Transcribe the full content of the PDF into plain text. Be as accurate as possible.
    2. Translate the transcribed text into Spanish.
    3. Provide a detailed quality assessment including:
       - Overall confidence score (0-100%)
       - Specific sections that may contain errors
       - Character recognition issues (garbled text, missing characters, etc.)
       - Formatting problems or layout issues

    Output the result in the following format:
    --- TRANSCRIPCIÓN ---
    [Original text here]

    --- TRADUCCIÓN ---
    [Spanish translation here]

    --- QUALITY ASSESSMENT ---
    Confidence Score: [X]%
    Issues Found: [List specific problems or "None identified"]
    Suspicious Sections: [Line numbers or text snippets that need manual review]
    Recommendations: [Any suggestions for improving accuracy]
    """

    response = model.generate_content([pdf_file, prompt])

    # Parse the response to extract transcription
    response_text = response.text
    transcription_start = response_text.find("--- TRANSCRIPCIÓN ---")
    translation_start = response_text.find("--- TRADUCCIÓN ---")

    if transcription_start != -1 and translation_start != -1:
        transcription = response_text[
            transcription_start + len("--- TRANSCRIPCIÓN ---") : translation_start
        ].strip()
    else:
        transcription = "Error: Could not extract transcription from response"

    # Detect anomalies in the transcription
    issues = detect_anomalies(transcription)
    confidence_score = calculate_confidence_score(transcription, issues)

    # Enhance the response with our analysis
    enhanced_response = response_text
    if issues:
        enhanced_response += f"\n\n--- AUTOMATED ANALYSIS ---\n"
        enhanced_response += f"Confidence Score: {confidence_score}%\n"
        enhanced_response += f"Detected Issues: {', '.join(issues)}\n"
        if confidence_score < 70:
            enhanced_response += "⚠️  LOW CONFIDENCE - Manual review recommended\n"
        elif issues:
            enhanced_response += "⚡ POTENTIAL ISSUES - Spot check recommended\n"
    else:
        enhanced_response += f"\n\n--- AUTOMATED ANALYSIS ---\n"
        enhanced_response += f"Confidence Score: {confidence_score}%\n"
        enhanced_response += "✅ No obvious issues detected\n"

    return enhanced_response



def save_to_file(text, output_path, output_format="docx"):
    """
    Saves the text to the specified path in the given format.
    Adds a mandatory AI disclaimer to the header/top of the document.
    """
    # Updated Disclaimer Text
    AI_DISCLAIMER_ES = (
        "Traducción no oficial.\n"
        "Realizado con servicios de traducción de google impulsado por IA.\n"
        "Puede contener errores"
    )
    
    # Enable multiple lines for PDF processing logic
    disclaimer_lines = AI_DISCLAIMER_ES.split('\n')

    # Ensure directory exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

    if output_format == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Add disclaimer to the header of the default section
        section = doc.sections[0]
        header = section.header
        
        # We want it to be right aligned, light gray
        for line in disclaimer_lines:
            paragraph = header.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(line)
            run.bold = False
            run.font.color.rgb = RGBColor(105, 105, 105) # Dim Gray (Darker)
            run.font.size = Pt(10) # Larger size

        # Add content
        doc.add_paragraph(text)
        doc.save(output_path)

    elif output_format == "pdf":
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.lib.utils import simpleSplit

        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Helper to draw header
        def draw_header(c):
            c.setFont("Helvetica", 10) # Larger size
            c.setFillColor(colors.dimgrey) # Darker gray
            
            # Position at top right
            # We'll calculate width of text to align right properly, or just use drawRightString
            y_pos = height - 20
            for line in disclaimer_lines:
                c.drawRightString(width - 40, y_pos, line)
                y_pos -= 12 # Increased spacing for larger font
            
            c.setFillColor(colors.black) # Reset

        lines = text.split('\n')
        
        # Draw header on first page
        draw_header(c)
        
        y = height - 60
        margin = 40
        line_height = 12
        
        c.setFont("Helvetica", 10)
        
        for line in lines:
            # Wrap line if too long
            wrapped_lines = simpleSplit(line, "Helvetica", 10, width - 2 * margin)
            for wrapped_line in wrapped_lines:
                if y < margin:
                    c.showPage()
                    draw_header(c)
                    c.setFont("Helvetica", 10)
                    y = height - 60
                
                c.drawString(margin, y, wrapped_line)
                y -= line_height
                
        c.save()

    else:
        raise ValueError("Unsupported format. Only 'docx' and 'pdf' are supported.")
